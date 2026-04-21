import os

def generate_samples():
    # Ensure a directory exists for samples
    sample_dir = "sample_documents"
    os.makedirs(sample_dir, exist_ok=True)
    
    # --- Sample 1: TXT format ---
    machine_learning_txt = """Introduction to Machine Learning
Machine learning is a subset of artificial intelligence that focuses on building systems that learn from data.
Instead of being explicitly programmed to perform a task, these systems improve their performance over time.

Supervised Learning
In supervised learning, algorithms are trained on labeled datasets. This means the data comes with an answer key.
Common applications include spam detection in emails and image recognition.

Unsupervised Learning
Unsupervised learning algorithms analyze unlabeled data to discover hidden patterns. 
Clustering is a famous technique used in customer segmentation and market basket analysis.
"""
    with open(os.path.join(sample_dir, "machine_learning_overview.txt"), "w", encoding="utf-8") as f:
        f.write(machine_learning_txt)
        
    # --- Sample 2: TXT format ---
    climate_change_txt = """Global Climate Change Report 2026
Climate change refers to long-term shifts in temperatures and weather patterns.
These shifts may be natural, but since the 1800s, human activities have been the main driver of climate change.
The primary cause is the burning of fossil fuels like coal, oil, and gas, which produces heat-trapping gases.

Impacts on Oceans
The ocean has absorbed a significant portion of the extra heat trapped in the climate system.
This causes thermal expansion and contributes to sea-level rise, threatening coastal communities worldwide.
Furthermore, increased carbon dioxide levels cause ocean acidification, severely impacting marine biodiversity, including coral reefs.
"""
    with open(os.path.join(sample_dir, "climate_change_report.txt"), "w", encoding="utf-8") as f:
        f.write(climate_change_txt)

    # --- Sample 3: DOCX format ---
    try:
        from docx import Document
        
        doc = Document()
        doc.add_heading('Financial Projection 2026', 0)
        
        doc.add_heading('Quarter 1 Overview', level=1)
        doc.add_paragraph('Our revenue for Q1 exceeded expectations by 15%. This was driven primarily by strong software sales in the European market.')
        
        doc.add_heading('Risk Factors', level=1)
        doc.add_paragraph('Inflationary pressures and supply chain disruptions continue to pose a risk to our hardware manufacturing division. We are actively seeking secondary suppliers to mitigate these vulnerabilities.')
        
        doc.save(os.path.join(sample_dir, "financial_projections.docx"))
        print("Generated DOCX sample successfully.")
    except ImportError:
        print("python-docx is not installed. Skipping DOCX generation.")
        print("To generate DOCX, run: pip install python-docx")

    print(f"Sample documents successfully generated in the '{sample_dir}' directory.")
    print("You can upload these files directly into the Streamlit app to test the semantic search!")

if __name__ == "__main__":
    generate_samples()
